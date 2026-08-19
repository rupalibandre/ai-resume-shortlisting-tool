from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def convert_docx_to_pdf(file_path: str) -> str:
    docx_path = Path(file_path)

    # Already a PDF
    if docx_path.suffix.lower() != ".docx":
        return str(docx_path)

    pdf_path = docx_path.with_suffix(".pdf")

    try:
        document = Document(str(docx_path))

        styles = getSampleStyleSheet()

        pdf = SimpleDocTemplate(
            str(pdf_path),
            pagesize=A4,
            rightMargin=15 * mm,
            leftMargin=15 * mm,
            topMargin=15 * mm,
            bottomMargin=15 * mm,
        )

        elements = []

        # Extract normal paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                elements.append(
                    Paragraph(
                        text.replace("&", "&amp;"),
                        styles["Normal"],
                    )
                )
                elements.append(Spacer(1, 5))

        # Extract tables too
        for table in document.tables:
            for row in table.rows:
                cells = []

                for cell in row.cells:
                    text = cell.text.strip()

                    if text:
                        cells.append(
                            text.replace("&", "&amp;")
                        )

                if cells:
                    row_text = " | ".join(cells)

                    elements.append(
                        Paragraph(
                            row_text,
                            styles["Normal"],
                        )
                    )
                    elements.append(Spacer(1, 4))

        if not elements:
            elements.append(
                Paragraph(
                    "No readable text found in DOCX file.",
                    styles["Normal"],
                )
            )

        pdf.build(elements)

    except Exception as e:
        raise Exception(
            f"DOCX to PDF conversion failed: {e}"
        )

    if not pdf_path.exists():
        raise Exception(
            f"PDF was not created: {pdf_path}"
        )

    return str(pdf_path)